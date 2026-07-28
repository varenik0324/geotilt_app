import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
import io
import re
import logging
import requests
import os
import json
import tempfile
from datetime import datetime
from typing import Optional, Tuple, Dict, Any, List

# ============================================================
# 1. КОНФИГУРАЦИЯ
# ============================================================
CONFIG = {
    "BOT_TOKEN": "8538186715:AAG7XsBxp6TAy2lalWQ6_KkBkrUIEZCqxuw",
    "CHAT_ID": "1278271780",
    "LOG_FILE": "app_errors.log",
    "PROFILES_FILE": "profiles.json",
    "DEFAULT_K_EM15H": 0.0031559,
    "DEFAULT_K_SM25H": 0.0035708,
    "F_STRING": 12.2,
    "F_CONCRETE": 10.0,
    "E_MODULUS": 3_000_000,
    "PILE_A": 6.51e-08,
    "PILE_B": -0.02931,
    "PILE_C": 248.4372,
    "PILE_K": -0.036375,
    "PILE_T_REF": 23.9,
}

# ============================================================
# 2. УТИЛИТЫ
# ============================================================
def clean_numeric(val):
    if pd.isna(val):
        return np.nan
    if isinstance(val, str):
        val = val.replace(',', '.').replace(' ', '').strip()
        return float(val) if val and val != '-' else np.nan
    return val

def send_telegram(message: str) -> bool:
    try:
        url = f"https://api.telegram.org/bot{CONFIG['BOT_TOKEN']}/sendMessage"
        payload = {"chat_id": CONFIG['CHAT_ID'], "text": f"📩 {message}", "parse_mode": "HTML"}
        r = requests.post(url, json=payload, timeout=5)
        return r.status_code == 200
    except Exception as e:
        logging.error(f"Telegram exception: {e}")
        return False

def save_profile(profile: Dict):
    profiles = {}
    if os.path.exists(CONFIG["PROFILES_FILE"]):
        with open(CONFIG["PROFILES_FILE"], 'r', encoding='utf-8') as f:
            profiles = json.load(f)
    profiles[profile['name']] = profile
    with open(CONFIG["PROFILES_FILE"], 'w', encoding='utf-8') as f:
        json.dump(profiles, f, ensure_ascii=False, indent=2)

def load_profiles() -> Dict:
    if os.path.exists(CONFIG["PROFILES_FILE"]):
        with open(CONFIG["PROFILES_FILE"], 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

# ============================================================
# 3. СПЕЦИФИКАЦИИ ДАТЧИКОВ (упрощённые)
# ============================================================
SENSOR_SPECS = {
    "MAS‑VWS‑EM15H (встроенный)": {"k_factor": "0.0031559"},
    "MAS‑VWS‑SM25H (поверхностный длинная база)": {"k_factor": "0.0035708"},
    "MAS‑VWS‑SM15 (поверхностный)": {"k_factor": "G × C"},
    "MAS‑VWE (давление грунта)": {"k_factor": "G × C"}
}

def get_sensor_specs(sensor_type: str) -> str:
    return f"Тип: {sensor_type}\nK: {SENSOR_SPECS.get(sensor_type, {}).get('k_factor', 'неизвестен')}"

# ============================================================
# 4. ОБРАБОТЧИК ТЕНЗОДАТЧИКОВ (ГАРАНТИРОВАННОЕ ПРИВЕДЕНИЕ К ЧИСЛАМ)
# ============================================================
class DataProcessor:
    @staticmethod
    def clean_and_convert(df: pd.DataFrame, col: str) -> pd.Series:
        if col not in df.columns:
            return pd.Series(index=df.index, dtype=float)
        series = df[col].astype(str).str.replace(',', '.').str.replace(' ', '').str.strip()
        series = series.replace('', np.nan)
        return pd.to_numeric(series, errors='coerce')

    @staticmethod
    def validate_data(df: pd.DataFrame) -> Tuple[bool, str, pd.DataFrame]:
        required = ['load', 'freq', 'temp']
        if df is None or df.empty:
            return False, "DataFrame пуст.", df
        missing = [c for c in required if c not in df.columns]
        if missing:
            return False, f"Отсутствуют столбцы: {', '.join(missing)}", df
        df_clean = df.copy()
        errors = []
        for col in required:
            converted = DataProcessor.clean_and_convert(df_clean, col)
            invalid_mask = converted.isna()
            if invalid_mask.any():
                invalid_rows = df_clean.index[invalid_mask].tolist()
                errors.append(f"В столбце '{col}' проблемы в строках: {invalid_rows[:10]}{'...' if len(invalid_rows)>10 else ''}")
            df_clean[col] = converted
        df_clean = df_clean.dropna(subset=required, how='all')
        for col in required:
            if df_clean[col].isna().sum() > 0:
                df_clean[col] = df_clean[col].interpolate(method='linear', limit=5)
        df_clean = df_clean.dropna(subset=required)
        if df_clean.empty:
            return False, "После очистки не осталось числовых строк. Проверьте данные.", df_clean
        # Принудительное приведение к float
        df_clean[required] = df_clean[required].astype(float)
        if errors:
            msg = "Обнаружены проблемы с данными:\n" + "\n".join(errors) + "\nПроблемные строки были удалены."
            return True, msg, df_clean
        else:
            return True, "Данные успешно проверены.", df_clean

    @staticmethod
    def process_strain_data(df: pd.DataFrame, f0: float, t0: float,
                            sensor_type: str, g_val: Optional[float] = None,
                            c_val: Optional[float] = None) -> Tuple[Optional[pd.DataFrame], Optional[Dict]]:
        # Полная защита от нечисловых данных
        if df is None or df.empty:
            return None, None

        df = df.copy()
        # Принудительное преобразование всех трёх столбцов в числа
        for col in ['load', 'freq', 'temp']:
            if col not in df.columns:
                return None, None
            df[col] = df[col].astype(str).str.replace(',', '.').str.replace(' ', '').str.strip()
            df[col] = pd.to_numeric(df[col], errors='coerce')
        df = df.dropna(subset=['load', 'freq', 'temp'])
        if df.empty:
            return None, None
        # Явный перевод в float
        df[['load', 'freq', 'temp']] = df[['load', 'freq', 'temp']].astype(float)

        # Приводим начальные параметры к float
        f0 = float(f0)
        t0 = float(t0)

        # Определяем K
        K = {
            'MAS‑VWS‑EM15H (встроенный)': CONFIG["DEFAULT_K_EM15H"],
            'MAS‑VWS‑SM25H (поверхностный длинная база)': CONFIG["DEFAULT_K_SM25H"]
        }.get(sensor_type)
        if K is None and sensor_type in ['MAS‑VWS‑SM15 (поверхностный)', 'MAS‑VWE (давление грунта)']:
            if g_val is None or c_val is None:
                return None, None
            K = g_val * c_val
        if K is None:
            return None, None

        # Расчёт
        df['strain'] = K * (df['freq']**2 - f0**2) + (df['temp'] - t0) * (CONFIG["F_STRING"] - CONFIG["F_CONCRETE"])
        df['stress_MPa'] = CONFIG["E_MODULUS"] * df['strain'] / 1_000_000 * 0.00689476

        # Русские заголовки для вывода
        df['Прирост деформации, μϵ'] = df['strain']
        df['Напряжение, МПа'] = df['stress_MPa']

        stats = {
            'Количество точек': len(df),
            'Средняя деформация, μϵ': df['Прирост деформации, μϵ'].mean(),
            'Макс. деформация, μϵ': df['Прирост деформации, μϵ'].max(),
            'Мин. деформация, μϵ': df['Прирост деформации, μϵ'].min(),
            'Среднее напряжение, МПа': df['Напряжение, МПа'].mean(),
            'Макс. напряжение, МПа': df['Напряжение, МПа'].max(),
            'Мин. напряжение, МПа': df['Напряжение, МПа'].min(),
            'Std деформация, μϵ': df['Прирост деформации, μϵ'].std(),
            'Статистика': df['Прирост деформации, μϵ'].describe().to_dict()
        }
        return df, stats

# ============================================================
# 5. ГЕНЕРАЦИЯ ОТЧЁТОВ (с русскими заголовками)
# ============================================================
class ReportGenerator:
    @staticmethod
    def to_excel(df: pd.DataFrame, stats: Dict, sensor_name: str, sensor_type: str) -> bytes:
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            out_df = df[['load', 'freq', 'temp', 'Прирост деформации, μϵ', 'Напряжение, МПа']].copy()
            out_df.rename(columns={
                'load': 'Нагрузка, тс',
                'freq': 'Частота, Гц',
                'temp': 'Температура, °C'
            }, inplace=True)
            out_df.to_excel(writer, index=False, sheet_name='Результат')
            stats_df = pd.DataFrame.from_dict(stats, orient='index', columns=['Значение'])
            stats_df.to_excel(writer, sheet_name='Сводка')
            ws_spec = writer.book.add_worksheet('Спецификация датчика')
            ws_spec.write(0, 0, get_sensor_specs(sensor_type))
        return output.getvalue()

    @staticmethod
    def to_pdf(df: pd.DataFrame, stats: Dict, sensor_name: str, sensor_type: str,
               f0: float, t0: float) -> io.BytesIO:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        import matplotlib.pyplot as plt
        from PIL import Image

        fig, ax = plt.subplots(figsize=(8, 4))
        ax.plot(df['load'], df['Прирост деформации, μϵ'], 'o-', color='#1f77b4', linewidth=2, markersize=8)
        ax.set_xlabel("Нагрузка, тс")
        ax.set_ylabel("Прирост деформации, μϵ")
        ax.set_title("Прирост деформации от нагрузки")
        ax.grid(True)
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=100)
        buf.seek(0)
        img = Image.open(buf)
        plt.close(fig)

        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, f"Отчёт по датчику: {sensor_name}")
        c.setFont("Helvetica", 12)
        c.drawString(50, height - 80, f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        c.drawString(50, height - 100, f"f₀ = {f0:.1f} Гц, T₀ = {t0:.1f} °C")
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, height - 130, "Спецификация датчика:")
        c.setFont("Helvetica", 10)
        c.drawString(55, height - 150, get_sensor_specs(sensor_type))
        img_path = tempfile.mktemp(suffix=".png")
        img.save(img_path)
        c.drawImage(img_path, 50, height - 450, width=500, height=250)
        os.remove(img_path)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, height - 480, "Сводка:")
        c.setFont("Helvetica", 10)
        y = height - 500
        for key, val in stats.items():
            if key not in ['Статистика']:
                c.drawString(60, y, f"{key}: {val:.3f}" if isinstance(val, float) else f"{key}: {val}")
                y -= 15
                if y < 50:
                    c.showPage()
                    y = height - 50
        c.save()
        buffer.seek(0)
        return buffer

    @staticmethod
    def to_word(df: pd.DataFrame, stats: Dict, sensor_name: str, sensor_type: str,
                f0: float, t0: float) -> io.BytesIO:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import matplotlib.pyplot as plt
        from PIL import Image

        doc = Document()
        title = doc.add_heading(f"Отчёт по датчику: {sensor_name}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph(f"f₀ = {f0:.1f} Гц, T₀ = {t0:.1f} °C")
        doc.add_heading("Спецификация датчика", level=2)
        doc.add_paragraph(get_sensor_specs(sensor_type))
        doc.add_heading("Сводка", level=2)
        for key, val in stats.items():
            if key not in ['Статистика']:
                doc.add_paragraph(f"{key}: {val:.3f}" if isinstance(val, float) else f"{key}: {val}")

        fig, ax = plt.subplots(figsize=(8, 4))
        ax.plot(df['load'], df['Прирост деформации, μϵ'], 'o-', color='#1f77b4', linewidth=2, markersize=8)
        ax.set_xlabel("Нагрузка, тс")
        ax.set_ylabel("Прирост деформации, μϵ")
        ax.set_title("Прирост деформации от нагрузки")
        ax.grid(True)
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=100)
        buf.seek(0)
        plt.close(fig)
        img = Image.open(buf)
        img_path = tempfile.mktemp(suffix=".png")
        img.save(img_path)
        doc.add_picture(img_path, width=Inches(6))
        os.remove(img_path)

        doc.add_heading("Таблица результатов (первые 20 строк)", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Нагрузка, тс"
        hdr_cells[1].text = "Частота, Гц"
        hdr_cells[2].text = "Температура, °C"
        hdr_cells[3].text = "Прирост деформации, μϵ"
        hdr_cells[4].text = "Напряжение, МПа"
        for _, row in df.head(20).iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = f"{row['load']:.1f}"
            row_cells[1].text = f"{row['freq']:.1f}"
            row_cells[2].text = f"{row['temp']:.1f}"
            row_cells[3].text = f"{row['Прирост деформации, μϵ']:.3f}"
            row_cells[4].text = f"{row['Напряжение, МПа']:.3f}"
        doc.add_paragraph("© Геофундамент, 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

# ============================================================
# 6. UI-ФУНКЦИИ (с русскими заголовками)
# ============================================================
def display_flat_results(result: pd.DataFrame, stats: Dict, sensor_name: str, sensor_type: str):
    st.subheader("✅ Результат обработки")
    display_df = result[['load', 'freq', 'temp', 'Прирост деформации, μϵ', 'Напряжение, МПа']].copy()
    display_df.rename(columns={
        'load': 'Нагрузка, тс',
        'freq': 'Частота, Гц',
        'temp': 'Температура, °C'
    }, inplace=True)
    st.dataframe(display_df.style.format({
        'Нагрузка, тс': '{:.2f}',
        'Частота, Гц': '{:.1f}',
        'Температура, °C': '{:.1f}',
        'Прирост деформации, μϵ': '{:.5f}',
        'Напряжение, МПа': '{:.5f}'
    }))

    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=result['load'],
        y=result['Прирост деформации, μϵ'],
        mode='lines+markers',
        name='Прирост деформации, μϵ'
    ))
    fig.update_layout(
        title="Прирост деформации от нагрузки",
        xaxis_title="Нагрузка, тс",
        yaxis_title="Прирост деформации, μϵ",
        template=st.session_state.get('template', 'plotly_white')
    )
    st.plotly_chart(fig, use_container_width=True)

    with st.expander("📊 Расширенная статистика"):
        stats_df = pd.DataFrame.from_dict(stats, orient='index', columns=['Значение'])
        st.dataframe(stats_df)

    col1, col2, col3 = st.columns(3)
    with col1:
        excel_data = ReportGenerator.to_excel(result, stats, sensor_name, sensor_type)
        st.download_button(
            label="📊 Excel",
            data=excel_data,
            file_name=f"result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    with col2:
        pdf_data = ReportGenerator.to_pdf(result, stats, sensor_name, sensor_type,
                                         st.session_state.get('f0', 1000.0),
                                         st.session_state.get('t0', 20.0))
        st.download_button(
            label="📄 PDF",
            data=pdf_data.getvalue(),
            file_name=f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
            mime="application/pdf"
        )
    with col3:
        word_data = ReportGenerator.to_word(result, stats, sensor_name, sensor_type,
                                           st.session_state.get('f0', 1000.0),
                                           st.session_state.get('t0', 20.0))
        st.download_button(
            label="📝 Word",
            data=word_data.getvalue(),
            file_name=f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ============================================================
# 7. ОСНОВНОЕ ПРИЛОЖЕНИЕ
# ============================================================
def main():
    st.set_page_config(
        page_title="Анализ датчиков | Геофундамент",
        page_icon="📐",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # Инициализация состояния
    for key in ['result', 'stats', 'sensor_name', 'template', 'f0', 't0', 'profiles', 'page']:
        if key not in st.session_state:
            if key == 'template':
                st.session_state[key] = 'plotly_white'
            elif key in ['f0', 't0']:
                st.session_state[key] = 1000.0 if key == 'f0' else 20.0
            elif key == 'page':
                st.session_state[key] = 'Главная'
            else:
                st.session_state[key] = None if key not in ['profiles'] else {}

    # Шапка
    col1, col2, col3 = st.columns([1, 3, 1])
    with col1:
        st.image("https://via.placeholder.com/80x40?text=GF", width=80)
    with col2:
        st.title("📊 Анализ данных тензодатчиков")
    with col3:
        theme_toggle = st.toggle("🌙 Тёмная тема", value=st.session_state.template == 'plotly_dark')
        st.session_state.template = 'plotly_dark' if theme_toggle else 'plotly_white'

    # Боковая панель
    with st.sidebar:
        st.markdown("### 🧭 Навигация")
        pages = {
            "🏠 Главная": "Главная",
            "📂 Загрузка": "Загрузка",
            "✏️ Ручной ввод": "Ручной ввод",
            "📚 Справка": "Справка"
        }
        selected_page = st.radio("", list(pages.keys()), index=list(pages.values()).index(st.session_state.page))
        st.session_state.page = pages[selected_page]

        st.markdown("---")
        st.markdown("### ⚙️ Параметры датчика")
        sensor_type = st.selectbox("Тип датчика", list(SENSOR_SPECS.keys()), key="sensor_type")
        specs = SENSOR_SPECS.get(sensor_type)
        if specs:
            st.caption(f"**K:** {specs.get('k_factor')}")
        g_val = c_val = None
        if sensor_type in ["MAS‑VWS‑SM15 (поверхностный)", "MAS‑VWE (давление грунта)"]:
            g_val = st.number_input("G", value=1.0, step=0.001, format="%.3f", key="g_val")
            c_val = st.number_input("C", value=1.0, step=0.001, format="%.3f", key="c_val")

        st.info("f₀ и T₀ автоматически определяются из первой строки данных.")
        st.caption(f"Текущие значения по умолчанию: f₀ = {st.session_state.f0:.1f} Гц, T₀ = {st.session_state.t0:.1f} °C (не используются)")

        st.markdown("---")
        profile_name = st.text_input("💾 Сохранить профиль", value="default")
        if st.button("💾 Сохранить"):
            profile = {
                'name': profile_name,
                'sensor_type': sensor_type,
                'f0': st.session_state.f0,
                't0': st.session_state.t0,
                'g_val': g_val,
                'c_val': c_val,
                'theme': 'Тёмная' if theme_toggle else 'Светлая'
            }
            save_profile(profile)
            st.success("Профиль сохранён!")

        profiles = load_profiles()
        if profiles:
            profile_names = list(profiles.keys())
            selected_profile = st.selectbox("📂 Загрузить профиль", [""] + profile_names)
            if selected_profile and st.button("📂 Загрузить"):
                p = profiles[selected_profile]
                st.session_state.sensor_type = p.get('sensor_type', 'MAS‑VWS‑EM15H (встроенный)')
                st.session_state.f0 = p.get('f0', 1000.0)
                st.session_state.t0 = p.get('t0', 20.0)
                if 'g_val' in p:
                    st.session_state.g_val = p['g_val']
                if 'c_val' in p:
                    st.session_state.c_val = p['c_val']
                st.rerun()

    page = st.session_state.page

    if page == "Главная":
        st.markdown("## 🏠 Дашборд")
        st.markdown("Добро пожаловать в приложение для анализа данных тензодатчиков!")
        if st.session_state.result is not None:
            st.markdown("### 📊 Последние результаты")
            res = st.session_state.result[['load', 'freq', 'temp', 'Прирост деформации, μϵ', 'Напряжение, МПа']].copy()
            res.rename(columns={
                'load': 'Нагрузка, тс',
                'freq': 'Частота, Гц',
                'temp': 'Температура, °C'
            }, inplace=True)
            st.dataframe(res.head(10))
        else:
            st.info("Нет загруженных данных. Перейдите в раздел 'Загрузка' или 'Ручной ввод'.")

    elif page == "Загрузка":
        st.markdown("## 📂 Загрузка файла (плоская таблица)")
        st.markdown("Файл должен содержать колонки: **нагрузка (load)**, **частота (freq)**, **температура (temp)**.")
        st.markdown("⚠️ **f₀ и T₀ будут автоматически взяты из ПЕРВОЙ строки данных (ступень 0).**")
        uploaded = st.file_uploader("Выберите Excel-файл", type=["xlsx", "xls"], key="flat_upload")
        if uploaded:
            try:
                df_raw = pd.read_excel(uploaded)
                st.write("📋 Предпросмотр загруженных данных:")
                st.dataframe(df_raw.head(10))

                if len(df_raw.columns) >= 3:
                    cols = df_raw.columns.tolist()
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        load_col = st.selectbox("Столбец с нагрузкой", cols, index=0)
                    with col2:
                        freq_col = st.selectbox("Столбец с частотой", cols, index=1 if len(cols)>1 else 0)
                    with col3:
                        temp_col = st.selectbox("Столбец с температурой", cols, index=2 if len(cols)>2 else 0)

                    if load_col and freq_col and temp_col:
                        df_mapped = df_raw[[load_col, freq_col, temp_col]].copy()
                        df_mapped.columns = ['load', 'freq', 'temp']

                        # Всегда берём из первой строки
                        f0_auto = df_mapped['freq'].iloc[0]
                        t0_auto = df_mapped['temp'].iloc[0]
                        st.info(f"Автоопределены: f₀ = {f0_auto:.1f} Гц, T₀ = {t0_auto:.1f} °C")

                        st.subheader("✏️ Редактирование данных (опционально)")
                        edited_df = st.data_editor(df_mapped, num_rows="dynamic", use_container_width=True)

                        if st.button("🚀 Обработать данные", key="process_flat"):
                            ok, msg, df_clean = DataProcessor.validate_data(edited_df)
                            if ok:
                                st.success(msg)
                                result, stats = DataProcessor.process_strain_data(
                                    df_clean, f0_auto, t0_auto, sensor_type, g_val, c_val
                                )
                                if result is not None and not result.empty:
                                    st.session_state.result = result
                                    st.session_state.stats = stats
                                    st.session_state.sensor_name = uploaded.name
                                    display_flat_results(result, stats, uploaded.name, sensor_type)
                                else:
                                    st.error("Ошибка расчёта. Проверьте данные и настройки датчика.")
                            else:
                                st.error(msg)
                else:
                    st.warning("Файл должен содержать минимум 3 колонки.")
            except Exception as e:
                st.error(f"Ошибка: {e}")
                logging.error(f"Ошибка в загрузке: {e}")

    elif page == "Ручной ввод":
        st.markdown("## ✏️ Ручной ввод данных")
        st.markdown("Вставьте данные в формате: **нагрузка, частота, температура**.")
        st.markdown("⚠️ **f₀ и T₀ будут автоматически взяты из ПЕРВОЙ строки введённых данных.**")
        st.markdown("Поддерживаются разделители: **запятая**, **табуляция**, **пробел**, **точка с запятой**.")

        delimiter = st.selectbox("Выберите разделитель",
                                 ["Авто", "Запятая (,)", "Табуляция (\\t)", "Пробел", "Точка с запятой (;)"],
                                 index=0)
        sep_map = {
            "Авто": None,
            "Запятая (,)": ",",
            "Табуляция (\\t)": "\t",
            "Пробел": " ",
            "Точка с запятой (;)": ";"
        }
        sep = sep_map[delimiter]

        text = st.text_area("Введите данные (каждая строка – одна точка)", height=200)

        if st.button("🔄 Предпросмотр", key="preview_manual"):
            if not text.strip():
                st.warning("Введите данные.")
            else:
                try:
                    lines = [line.strip() for line in text.strip().splitlines() if line.strip()]
                    rows = []
                    for line in lines:
                        if sep:
                            parts = line.split(sep)
                        else:
                            parts = re.split(r'[,\t; ]+', line)
                        parts = [p.strip() for p in parts if p.strip()]
                        if len(parts) >= 3:
                            rows.append(parts[:3])
                    if not rows:
                        st.error("Не удалось распознать данные. Проверьте разделитель.")
                    else:
                        df_preview = pd.DataFrame(rows, columns=['load', 'freq', 'temp'])
                        st.write("📋 Распознанные данные:")
                        st.dataframe(df_preview)
                        st.session_state['manual_df'] = df_preview
                except Exception as e:
                    st.error(f"Ошибка предпросмотра: {e}")

        if 'manual_df' in st.session_state and st.session_state['manual_df'] is not None:
            st.subheader("✏️ Редактирование данных")
            edited_df = st.data_editor(st.session_state['manual_df'], num_rows="dynamic", use_container_width=True)

            if st.button("🚀 Обработать данные", key="process_manual"):
                f0_auto = edited_df['freq'].iloc[0]
                t0_auto = edited_df['temp'].iloc[0]
                ok, msg, df_clean = DataProcessor.validate_data(edited_df)
                if ok:
                    st.success(msg)
                    result, stats = DataProcessor.process_strain_data(
                        df_clean, f0_auto, t0_auto, sensor_type, g_val, c_val
                    )
                    if result is not None and not result.empty:
                        st.session_state.result = result
                        st.session_state.stats = stats
                        st.session_state.sensor_name = "Ручной ввод"
                        display_flat_results(result, stats, "Ручной ввод", sensor_type)
                    else:
                        st.error("Ошибка расчёта.")
                else:
                    st.error(msg)

    elif page == "Справка":
        st.markdown("## 📚 Справка")
        st.markdown("""
        **Приложение для расчёта деформации по данным тензодатчиков.**
        - Загрузите Excel-файл с колонками: нагрузка, частота, температура.
        - Либо введите данные вручную.
        - Укажите параметры датчика (тип, K и т.д.).
        - **f₀ и T₀ автоматически берутся из первой строки данных.**
        - Получите график прироста деформации от нагрузки и статистику.
        - Экспортируйте результаты в Excel, PDF или Word.
        """)

if __name__ == "__main__":
    logging.basicConfig(filename=CONFIG["LOG_FILE"], level=logging.INFO)
    main()
